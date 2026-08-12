# backend/tests/test_document_parser.py
import io

from docx import Document as DocxDocument

from app.services.document_parser import parse_text, split_chunks


def _make_docx_with_table():
    """构造含段落 + 数据表格的 docx（营销活动方案/报表形态）。"""
    doc = DocxDocument()
    doc.add_paragraph("国庆营销活动复盘方案")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "GMV"
    table.cell(0, 1).text = "860万"
    table.cell(1, 0).text = "ROI"
    table.cell(1, 1).text = "5.2"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_parse_docx_includes_tables():
    """docx 解析必须同时提取段落与表格（方案+数据表格混合文档，丢表格即丢核心数据）。"""
    text = parse_text(_make_docx_with_table(), "docx")
    assert "国庆营销活动复盘方案" in text  # 段落
    assert "GMV" in text and "860万" in text  # 表格内容
    assert "ROI" in text and "5.2" in text
    assert "GMV | 860万" in text

def test_split_chunks_basic():
    text = "段落一。" * 200
    chunks = split_chunks(text, chunk_size=100)
    assert len(chunks) > 1
    assert "段落一" in chunks[0]

def test_short_text_not_split():
    chunks = split_chunks("简短文本", chunk_size=500)
    assert len(chunks) == 1

def test_parse_markdown():
    # parse_text 接收 bytes 内容
    assert "标题" in parse_text("# 标题\n正文".encode(), "md")

def test_markdown_header_split():
    text = "# 第一章\n## 第一节\n内容A\n## 第二节\n内容B"
    chunks = split_chunks(text, ext="md")
    assert len(chunks) >= 2
    assert "第一章" in chunks[0]
