# backend/tests/test_experience_upload.py
import io
from datetime import date

import pytest
from docx import Document as DocxDocument
from httpx import AsyncClient, ASGITransport
from sqlalchemy import func, select

from app.main import app
from app.models.experience import Experience
from app.services.experience_svc import DistillOutput


async def _register(transport, username):
    """注册并登录，返回 headers。"""
    async with AsyncClient(transport=transport, base_url="http://test") as c:
        await c.post("/api/auth/register", json={"username": username, "password": "x123456", "display_name": "U"})
        r = await c.post("/api/auth/login", json={"username": username, "password": "x123456"})
        return {"Authorization": f"Bearer {r.json()['access_token']}"}


@pytest.mark.asyncio
async def test_upload_markdown_creates_draft(db_session, monkeypatch):
    """上传营销活动 md 文件 → 解析 + LLM 提炼 → 落库个人草稿经验。"""
    class FakeLLM:
        def with_structured_output(self, schema):
            return self
        async def ainvoke(self, prompt):
            assert "国庆营销" in prompt  # 文件内容确实喂给了 LLM
            return DistillOutput(
                title="国庆大促复盘", summary="满减+直播，ROI 5",
                content="详情", tags=["营销"],
                event_time=date(2024, 10, 1), result_metrics={"gmv": 320, "roi": 5},
            )
    monkeypatch.setattr("app.services.experience_svc.ModelFactory.get_llm", lambda: FakeLLM())
    async def _embed(t):
        return [[0.1] * 1536]
    monkeypatch.setattr("app.services.experience_svc.embed_texts", _embed)

    transport = ASGITransport(app=app)
    h = await _register(transport, "up1")
    async with AsyncClient(transport=transport, base_url="http://test") as c:
        files = {"file": ("活动复盘.md", "# 国庆营销活动\n国庆大促，满减+直播，GMV 320 万，ROI 5".encode(),
                          "text/markdown")}
        r = await c.post("/api/experiences/upload", files=files, headers=h)
        assert r.status_code == 200, r.text
        assert r.json()["title"] == "国庆大促复盘"
    # 落库为个人草稿经验，event_time/result_metrics 带出
    exp = (await db_session.scalars(select(Experience))).first()
    assert exp is not None
    assert exp.status == "draft" and exp.scope == "personal"
    assert exp.title == "国庆大促复盘"
    assert exp.event_time == date(2024, 10, 1)
    assert exp.result_metrics == {"gmv": 320, "roi": 5}


@pytest.mark.asyncio
async def test_upload_docx_with_table_includes_table_data(db_session, monkeypatch):
    """上传方案+数据表格混合 docx：表格内容（效果指标）必须进 LLM prompt。

    真实事故：parse_text 只提段落不提表格，营销活动 docx 的效果指标全丢，
    LLM 因缺 result_metrics 判无价值 → 400。"""
    class FakeLLM:
        def with_structured_output(self, schema):
            return self
        async def ainvoke(self, prompt):
            assert "GMV" in prompt and "860万" in prompt  # 表格数据确实喂给了 LLM
            return DistillOutput(
                title="国庆大促复盘", summary="直播+满减，ROI 5.2",
                content="详情", tags=["营销"],
                event_time=date(2024, 10, 1), result_metrics={"gmv": 860, "roi": 5.2},
            )
    monkeypatch.setattr("app.services.experience_svc.ModelFactory.get_llm", lambda: FakeLLM())
    async def _embed(t):
        return [[0.1] * 1536]
    monkeypatch.setattr("app.services.experience_svc.embed_texts", _embed)

    doc = DocxDocument()
    doc.add_paragraph("国庆营销活动复盘方案")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "GMV"
    table.cell(0, 1).text = "860万"
    table.cell(1, 0).text = "ROI"
    table.cell(1, 1).text = "5.2"
    buf = io.BytesIO()
    doc.save(buf)

    transport = ASGITransport(app=app)
    h = await _register(transport, "up4")
    async with AsyncClient(transport=transport, base_url="http://test") as c:
        files = {"file": ("活动复盘.docx", buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        r = await c.post("/api/experiences/upload", files=files, headers=h)
        assert r.status_code == 200, r.text
        assert r.json()["title"] == "国庆大促复盘"


@pytest.mark.asyncio
async def test_upload_unvaluable_returns_400(db_session, monkeypatch):
    """上传内容不构成营销经验 → 400 且不落库。"""
    class FakeLLM:
        def with_structured_output(self, schema):
            return self
        async def ainvoke(self, prompt):
            return DistillOutput(title=None, summary="", content="")
    monkeypatch.setattr("app.services.experience_svc.ModelFactory.get_llm", lambda: FakeLLM())
    async def _embed(t):
        return [[0.1] * 1536]
    monkeypatch.setattr("app.services.experience_svc.embed_texts", _embed)

    transport = ASGITransport(app=app)
    h = await _register(transport, "up2")
    async with AsyncClient(transport=transport, base_url="http://test") as c:
        files = {"file": ("普通.txt", "今天天气不错".encode(), "text/plain")}
        r = await c.post("/api/experiences/upload", files=files, headers=h)
        assert r.status_code == 400
    assert (await db_session.scalar(select(func.count()).select_from(Experience))) == 0


@pytest.mark.asyncio
async def test_upload_title_null_string_returns_400(db_session, monkeypatch):
    """LLM 把 JSON null 输出成字符串 'null' → 视为无价值，400 不落库。

    真实事故：title='null' 字符串被 if not title 判为真值，落库了内容全空的无效经验。"""
    class FakeLLM:
        def with_structured_output(self, schema):
            return self
        async def ainvoke(self, prompt):
            return DistillOutput(title="null", summary="", content="")
    monkeypatch.setattr("app.services.experience_svc.ModelFactory.get_llm", lambda: FakeLLM())
    async def _embed(t):
        return [[0.1] * 1536]
    monkeypatch.setattr("app.services.experience_svc.embed_texts", _embed)

    transport = ASGITransport(app=app)
    h = await _register(transport, "up5")
    async with AsyncClient(transport=transport, base_url="http://test") as c:
        files = {"file": ("无价值.txt", "普通文本".encode(), "text/plain")}
        r = await c.post("/api/experiences/upload", files=files, headers=h)
        assert r.status_code == 400
    assert (await db_session.scalar(select(func.count()).select_from(Experience))) == 0


@pytest.mark.asyncio
async def test_upload_bad_pdf_returns_500(db_session):
    """损坏 pdf 文件 → 解析失败 → 500 且不落库。"""
    transport = ASGITransport(app=app)
    h = await _register(transport, "up3")
    async with AsyncClient(transport=transport, base_url="http://test") as c:
        files = {"file": ("bad.pdf", b"%PDF-1.4 broken garbage", "application/pdf")}
        r = await c.post("/api/experiences/upload", files=files, headers=h)
        assert r.status_code == 500
    assert (await db_session.scalar(select(func.count()).select_from(Experience))) == 0
