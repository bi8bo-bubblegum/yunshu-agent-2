import io

from pypdf import PdfReader
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter, MarkdownHeaderTextSplitter

def parse_text(content: bytes, ext: str) -> str:
    ext = ext.lower().lstrip(".")
    if ext == "pdf":
        reader = PdfReader(io.BytesIO(content))
        return "\n".join(p.extract_text() or "" for p in reader.pages)
    if ext in ["doc", "docx"]:
        doc = DocxDocument(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs]
        # 方案/报表文档的核心数据（渠道/预算/效果指标）往往在表格里，
        # 只提段落会丢失表格数据（真实事故：营销活动 docx 提炼不出经验）。
        for t in doc.tables:
            for row in t.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells))
        return "\n".join(parts)
    return content.decode("utf-8", errors="ignore")

def split_chunks(text: str, chunk_size: int = 500, overlap: int = 50, ext: str = "txt") -> list[str]:
    if ext in ["md", "markdown"]:
        splitter = MarkdownHeaderTextSplitter(headers_to_split_on=[("#","H1"),("##","H2"),("###","H3")])
        docs = splitter.split_text(text)
        return [
            (" ".join(
                f"{k}: {v}" for k, v in d.metadata.items()) + "\n" + d.page_content) if d.metadata else d.page_content
            for d in docs
        ]
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=overlap,
        separators=["\n\n", "\n", "。", "！", "？", "；", " ", ""],
    )
    return splitter.split_text(text)
